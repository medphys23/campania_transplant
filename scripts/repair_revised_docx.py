"""
Repair the revised Campania transplant manuscript:
- correct article-backed table values and labels
- recreate missing figures 12.2, 13.1, and 14.1
- restore caption/table structure in Sections 13 and 17
- emit an audit summary of supported vs unsupported items
"""
from __future__ import annotations

import io
import math
import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.document import Document as _Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
DOC_PATH = ROOT / "documents" / "campania_transplant_final_revised.docx"
OUT_PATH = ROOT / "documents" / "campania_transplant_final_repaired.docx"
FIG_DIR = ROOT / "documents" / "image"
AUDIT_PATH = ROOT / "documents" / "REVISED_VALIDATION_SUMMARY.md"


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = getattr(parent, "element", parent)
        if hasattr(parent_elm, "body"):
            parent_elm = parent_elm.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def normalize(text: str) -> str:
    return " ".join((text or "").split())


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def insert_paragraph_after_table(table: Table) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    return Paragraph(new_p, table._parent)


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
        paragraph._p = paragraph._element = None  # type: ignore[assignment]


def move_table_after_paragraph(table: Table, paragraph: Paragraph) -> None:
    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)
    paragraph._p.addnext(tbl)


def build_figure_12_2(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 6.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    def box(x, y, w, h, text, fc="#eef3f8", ec="#4a5568", size=12, weight="normal"):
        rect = plt.Rectangle((x, y), w, h, facecolor=fc, edgecolor=ec, linewidth=1.8)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=size, weight=weight, wrap=True)

    def arrow(x1, y1, x2, y2):
        ax.annotate(
            "",
            xy=(x2, y2),
            xytext=(x1, y1),
            arrowprops=dict(arrowstyle="->", lw=1.8, color="#2d3748"),
        )

    box(0.37, 0.82, 0.26, 0.1, "Total Annual Savings\n(S_TOTAL)", fc="#d9ecff", weight="bold", size=13)
    box(0.05, 0.50, 0.25, 0.18, "Module 1\nTransplant Expansion\nS_TX(t)\nBase case: ~€19.32M/year", fc="#e6f4ea")
    box(0.375, 0.50, 0.25, 0.18, "Module 2\nPIRP Incidence Reduction\nS_PIRP(t)\nBase case: ~€6.0M/year", fc="#fff4d6")
    box(0.70, 0.50, 0.25, 0.18, "Module 3\nPre-emptive Add-on\nS_PRE(t)\nBase case: ~€0.48M/year", fc="#f7e6ff")

    box(0.05, 0.18, 0.25, 0.16, "Inputs\nTx target 35 pmp\nTx baseline 14.3 pmp\nc_dial €50k\nc_tx1 €60k\nc_txm €12k", size=10.5)
    box(0.375, 0.18, 0.25, 0.16, "Inputs\nPop 5.8M\nIncidence 200 pmp/year\nPIRP reduction 10%\nAvoidInc = 116/year", size=10.5)
    box(0.70, 0.18, 0.25, 0.16, "Inputs\nAddTx ≈ 120/year\nPre-emptive share 20%\nΔ_pre €20k", size=10.5)

    arrow(0.175, 0.68, 0.5, 0.82)
    arrow(0.5, 0.68, 0.5, 0.82)
    arrow(0.825, 0.68, 0.5, 0.82)
    arrow(0.175, 0.34, 0.175, 0.50)
    arrow(0.5, 0.34, 0.5, 0.50)
    arrow(0.825, 0.34, 0.825, 0.50)

    ax.text(
        0.5,
        0.05,
        "Three mutually exclusive modules combine to estimate planning-level annual savings under the proposal's base-case assumptions.",
        ha="center",
        va="center",
        fontsize=10.5,
        color="#334155",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_figure_13_1(path: Path) -> None:
    years = [1, 2, 3, 5, 10]
    dialysis = [50000, 100000, 150000, 250000, 500000]
    transplant = [60000, 72000, 84000, 108000, 168000]

    fig, ax = plt.subplots(figsize=(10.5, 6.2))
    ax.plot(years, dialysis, marker="o", linewidth=2.4, color="#b91c1c", label="Dialysis")
    ax.plot(years, transplant, marker="o", linewidth=2.4, color="#0f766e", label="Kidney transplantation")
    ax.axvline(2, color="#475569", linestyle="--", linewidth=1.5)
    ax.annotate("Break-even\n~Year 2", xy=(2, 72000), xytext=(2.6, 175000), arrowprops=dict(arrowstyle="->", lw=1.2))
    ax.set_title("Ten-year cumulative cost trajectories per patient")
    ax.set_xlabel("Year")
    ax.set_ylabel("Cumulative cost (EUR)")
    ax.set_xticks(years)
    ax.set_ylim(0, 540000)
    ax.grid(axis="y", alpha=0.25)
    ax.legend(frameon=False)
    ax.text(
        0.02,
        0.02,
        "Base case: c_dial=€50k, c_tx1=€60k, c_txm=€12k",
        transform=ax.transAxes,
        fontsize=10,
        color="#334155",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def build_figure_14_1(path: Path) -> None:
    dialysis_patients = 6500
    dialysis_cost = 50000
    transplant_patients = (2081 + 2101 + 2240 + 2245) / 4
    transplant_maint_cost = 12000

    dialysis_total = dialysis_patients * dialysis_cost / 1_000_000
    transplant_total = transplant_patients * transplant_maint_cost / 1_000_000

    fig, ax = plt.subplots(figsize=(9.6, 6.0))
    labels = ["Maintenance dialysis", "Transplant maintenance"]
    values = [dialysis_total, transplant_total]
    colors = ["#c2410c", "#2563eb"]
    bars = ax.bar(labels, values, color=colors, width=0.55)
    ax.set_ylabel("Approximate annual cost (EUR millions)")
    ax.set_title("Current annual ESKD cost distribution in Campania")
    ax.set_ylim(0, max(values) * 1.22)
    ax.grid(axis="y", alpha=0.2)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 4, f"€{value:.1f}M", ha="center", va="bottom", fontsize=11)
    ax.text(
        0.5,
        -0.16,
        "Approximate base-case breakdown using ~6,500 dialysis patients at €50k/year and Campania transplant prevalence\n(2015–2018 midpoint) with €12k annual post-transplant maintenance cost.",
        transform=ax.transAxes,
        ha="center",
        va="top",
        fontsize=9.5,
        color="#334155",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def find_block_indices(doc: Document) -> dict[str, int]:
    targets = {
        "fig12_2": "Figure 12.2: Corrected economic model structure. Three mutually exclusive modules contribute to total savings.",
        "fig13_1": "Figure 13.1: Ten-year cumulative cost trajectories for dialysis and kidney transplantation per patient.",
        "fig14_1": "Figure 14.1: Current annual ESKD cost distribution in Campania.",
        "table17_1": "Table 17.1. Aggregate annual economic impact.",
        "table17_2_sentence": "Table 17.2 summarizes the annual savings components and the aggregate projected impact.",
        "section13_2": "13.2 Pre-Emptive Transplantation: Incremental Economic Benefit",
        "section14_2": "14.2 Scenario Analysis: Implementation of Emilia-Romagna–Level Transplant Activity",
    }
    found: dict[str, int] = {}
    for idx, block in enumerate(iter_block_items(doc)):
        if isinstance(block, Paragraph):
            text = normalize(block.text)
            for key, target in targets.items():
                if text == target:
                    found[key] = idx
    return found


def add_picture_after(paragraph: Paragraph, image_path: Path, width_in: float = 6.6) -> Paragraph:
    pic_p = paragraph_after(paragraph)
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    pic_p.alignment = 1
    return pic_p


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def write_audit_summary() -> None:
    AUDIT_PATH.write_text(
        """# Revised Manuscript Validation Summary

## Article-backed corrections applied

- `Table 4.1`: corrected Campania kidney transplant prevalence row to `2,081–2,245` patients and `355–389 pmp` based on Cirillo et al. `[7]`.
- `Table 4.2`: corrected annual kidney transplantation incidence row to `159–191` cases and `27–33 pmp` based on Cirillo et al. `[7]`.
- `Table 9.1`: corrected CT-PIRP subgroup counts to `230, 378, 152, 264, 90, 410, 741` and updated percentages to align with Rucci et al. `[8]` / Gibertoni et al. `[9]`.
- `Table 9.1`: corrected subgroup definitions from `Age 67` to `Age <67` for subgroups 4 and 5.
- Abstract: corrected 3-year cumulative savings per transplant to `~€66,000`, matching the manuscript's own `Table 13.1` base-case series.

## Figures recreated

- `Figure 12.2`: redrawn from the manuscript's described three-module economic model and Appendix X base-case assumptions.
- `Figure 13.1`: recreated from `Table 13.1` cumulative cost values.
- `Figure 14.1`: recreated as an approximate base-case cost distribution using explicit manuscript assumptions plus Campania transplant prevalence from Cirillo et al. `[7]`.

## Structural repairs applied

- Reinserted missing figure objects after captions for `Figures 12.2`, `13.1`, and `14.1`.
- Restored `Table 13.1` as a proper standalone caption above the cumulative-cost table rather than as narrative text.
- Restored `Table 17.2` as a proper standalone caption above the annual-impact table.
- Normalized the wording of the `CKD progression reduction (PIRP)` row in `Table 17.2`.

## Unsupported or partially supported items

- References `[13]` and `[16]` are still not present in the `articles` folder and could not be validated directly.
- Sections `12–17` remain partly dependent on internal model assumptions rather than direct article evidence.
- `Figure 14.1` is necessarily partly manuscript-derived because the `articles` folder does not include raw economic spreadsheets or original chart files.
""",
        encoding="utf-8",
    )


def main() -> int:
    if hasattr(sys.stdout, "buffer"):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig12 = FIG_DIR / "figure_12_2_recreated.png"
    fig13 = FIG_DIR / "figure_13_1_recreated.png"
    fig14 = FIG_DIR / "figure_14_1_recreated.png"
    build_figure_12_2(fig12)
    build_figure_13_1(fig13)
    build_figure_14_1(fig14)

    doc = Document(DOC_PATH)

    # Article-backed table fixes.
    table_4_1 = doc.tables[0]
    table_4_1.rows[3].cells[1].text = "2,081–2,245"
    table_4_1.rows[3].cells[2].text = "355–389 pmp"

    table_4_2 = doc.tables[1]
    table_4_2.rows[2].cells[1].text = "159–191"
    table_4_2.rows[2].cells[2].text = "27–33 pmp"

    table_9_1 = doc.tables[19]
    corrected_rows = [
        ("1", "Proteinuric, eGFR >33", "230 (10.1%)", "–3.66 (fastest)", "Low (short-term)", "Low", "HIGHEST"),
        ("2", "Proteinuric, eGFR ≤33, Phos ≤4.3", "378 (16.1%)", "–2.0 to –2.5", "High", "Moderate", "High"),
        ("3", "Proteinuric, eGFR ≤33, Phos >4.3", "152 (4.4%)", "–2.83", "Highest", "Highest", "HIGHEST"),
        ("4", "Non-proteinuric, Age <67, No DM", "264 (10.9%)", "–1.0 to –1.5", "Lowest", "Lowest", "Lowest"),
        ("5", "Non-proteinuric, Age <67, DM", "90 (3.9%)", "–1.5 to –2.0", "Moderate", "Moderate", "Moderate"),
        ("6", "Non-proteinuric, Age ≥67, Female", "410 (19.3%)", "–0.5 to –1.0 (stable)", "Very Low", "Moderate", "Low"),
        ("7", "Non-proteinuric, Age ≥67, Male", "741 (35.2%)", "–1.5 to –2.0", "Moderate", "High", "Moderate"),
    ]
    for row_idx, values in enumerate(corrected_rows, start=1):
        for cell_idx, value in enumerate(values):
            table_9_1.rows[row_idx].cells[cell_idx].text = value

    # Abstract correction.
    for p in doc.paragraphs:
        text = normalize(p.text)
        if text == "3-year cumulative savings per transplant: €36,000–€45,000":
            set_paragraph_text(p, "3-year cumulative savings per transplant: ~€66,000")
            break

    blocks = list(iter_block_items(doc))

    figure12_caption = next(b for b in blocks if isinstance(b, Paragraph) and normalize(b.text) == "Figure 12.2: Corrected economic model structure. Three mutually exclusive modules contribute to total savings.")
    add_picture_after(figure12_caption, fig12, width_in=6.8)

    # Rebuild Section 13 structure.
    blocks = list(iter_block_items(doc))
    fig13_caption = next(b for b in blocks if isinstance(b, Paragraph) and normalize(b.text) == "Figure 13.1: Ten-year cumulative cost trajectories for dialysis and kidney transplantation per patient.")
    table13 = next(b for b in blocks if isinstance(b, Table) and normalize(b.rows[0].cells[0].text) == "Year")
    table13_caption_para = next(
        b
        for b in blocks
        if isinstance(b, Paragraph)
        and normalize(b.text).startswith("Table 13.1 presents the per-patient cumulative cost comparison over 10 years.")
    )
    set_paragraph_text(table13_caption_para, "Table 13.1. Cumulative cost comparison per patient over 10 years.")
    move_table_after_paragraph(table13, table13_caption_para)
    after_table = insert_paragraph_after_table(table13)
    set_paragraph_text(
        after_table,
        "These values describe single-cohort savings for one transplant performed at year 0; program-level savings depend on annual transplant volume and are therefore modelled separately using the incident-flow framework described in Section 12.2.",
    )
    add_picture_after(fig13_caption, fig13, width_in=6.8)

    # Recreate Figure 14.1.
    blocks = list(iter_block_items(doc))
    fig14_caption = next(b for b in blocks if isinstance(b, Paragraph) and normalize(b.text) == "Figure 14.1: Current annual ESKD cost distribution in Campania.")
    add_picture_after(fig14_caption, fig14, width_in=6.6)

    # Table 17.2 caption and row text cleanup.
    for p in doc.paragraphs:
        text = normalize(p.text)
        if text == "Table 17.2 summarizes the annual savings components and the aggregate projected impact.":
            set_paragraph_text(p, "Table 17.2. Aggregate annual economic impact.")
        elif text == "CKD progression reduction (PIRP, )":
            set_paragraph_text(p, "CKD progression reduction (PIRP)")

    table17_2 = doc.tables[29]
    if normalize(table17_2.rows[3].cells[0].text) == "CKD progression reduction (PIRP, )":
        table17_2.rows[3].cells[0].text = "CKD progression reduction (PIRP)"

    # Keep Table 17.1 caption before its table if already present; ensure wording is unchanged.
    write_audit_summary()
    doc.save(OUT_PATH)
    print(f"Wrote repaired copy to {OUT_PATH}")
    print(f"Wrote figures to {FIG_DIR}")
    print(f"Wrote audit summary to {AUDIT_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
